"""Build the printable Muster teacher guide (.docx) for PD handouts / binders.

Same content as the in-app /help page (templates/help.html) -- if you change one,
change the other. The web page is the everyday link; this is the paper version.

    python build_teacher_guide.py

Writes C:\\Users\\there\\Downloads\\Muster_Teacher_Guide.docx
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = os.path.join(os.path.expanduser('~'), 'Downloads', 'Muster_Teacher_Guide.docx')

DARK = RGBColor(0x21, 0x25, 0x29)
BLUE = RGBColor(0x0D, 0x6E, 0xFD)
GREY = RGBColor(0x66, 0x66, 0x66)


def setup(doc):
    """Tight margins and a readable base font -- this gets printed and read at a desk."""
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.6)
        s.left_margin = s.right_margin = Inches(0.7)
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)


def title(doc):
    p = doc.add_paragraph()
    r = p.add_run('How to Use Muster')
    r.font.size, r.bold, r.font.color.rgb = Pt(24), True, DARK
    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    r = p.add_run('The hall pass app \u2014 teacher quick reference')
    r.font.size, r.font.color.rgb = Pt(11), GREY
    p.paragraph_format.space_after = Pt(10)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run('The one-sentence version: ')
    r.bold, r.font.size, r.font.color.rgb = True, Pt(11), BLUE
    r = p.add_run('every student has a QR code. Scan it (or tap their name), pick where '
                  "they're going, tap the blue button. Scan again when they come back.")
    r.font.size = Pt(11)


def h(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(13)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold, r.font.size, r.font.color.rgb = True, Pt(14), DARK
    # Rule under the heading, so sections are findable when flipping pages.
    pr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), '212529')
    bottom.set(qn('w:space'), '2')
    borders.append(bottom)
    pr.append(borders)


def sub(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold, r.font.size = True, Pt(11)


def steps(doc, items):
    """Numbered do-this-then-that list."""
    for i, t in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f'{i}.  ')
        r.bold = True
        _rich(p, t)


def bullets(doc, items, indent=0.28):
    for t in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(indent)
        p.paragraph_format.space_after = Pt(2)
        _rich(p, t)


def _rich(p, text):
    """**bold** segments inline, so button names stand out on paper."""
    for i, chunk in enumerate(text.split('**')):
        if chunk:
            r = p.add_run(chunk)
            r.bold = (i % 2 == 1)


def callout(doc, label, body, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(label + ' ')
    r.bold, r.font.color.rgb, r.font.size = True, color, Pt(10.5)
    _rich(p, body)


def code(doc, lines):
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(ln)
        r.font.name, r.font.size = 'Consolas', Pt(9.5)


def build():
    doc = Document()
    setup(doc)
    title(doc)

    # ── 1 ──────────────────────────────────────────────
    h(doc, '1.  Signing in')
    steps(doc, [
        'Open the Muster link on your computer or phone.',
        'Tap the red **Sign in with Google** button.',
        'Pick your **school Google account** \u2014 the same one you use for school email.',
        "That's it. You land on your roster page.",
    ])
    bullets(doc, [
        'You only have to do this once per device. It remembers you for about a month.',
        'Your **first sign-in is what creates your account**, so nothing looks set up until '
        "you've done it.",
        '**Sign out** is top-right, but you usually don\u2019t need it \u2014 stay signed in on '
        'your classroom computer.',
    ])
    callout(doc, 'If it says you\u2019re not authorized:',
            'your email hasn\u2019t been added to the approved list yet, or you signed in with a '
            'personal Gmail instead of your school account. Check which account you used, then '
            'ask Ms. Howe.', RGBColor(0xB5, 0x84, 0x00))

    # ── 2 ──────────────────────────────────────────────
    h(doc, '2.  Getting your students in')
    doc.add_paragraph('Two ways. Both live on the Roster page at the top.')
    sub(doc, 'One student at a time')
    steps(doc, [
        'Tap **+ Add Student**.',
        'Type first name, last name, pick a grade.',
        'Pick the **class period** they\u2019re in.',
        'Tap **Add to My Roster**.',
    ])
    sub(doc, 'A whole class at once')
    steps(doc, [
        'Tap **\u2191 Upload CSV**.',
        'Choose the **class period first** \u2014 everyone in the file goes into that period.',
        'Choose your file, tap **Upload**.',
        'You get a little report: how many added, how many were already there.',
    ])
    doc.add_paragraph('Your file needs exactly three columns, spelled like this:')
    code(doc, ['first_name,last_name,grade', 'Jane,Smith,9', 'John,Doe,11'])
    bullets(doc, [
        'Grades must be 7 through 12.',
        "**Use this year\u2019s grade**, not last year\u2019s. A 9th grader who\u2019s now in "
        '10th needs to say 10.',
        'Upload one file per class period. Six classes = six uploads.',
    ])
    callout(doc, 'Got a spreadsheet instead of a CSV?',
            'In Excel or Google Sheets choose File \u2192 Download / Save As \u2192 CSV. Make sure '
            'the very first row is first_name,last_name,grade')
    sub(doc, 'Fixing your roster')
    bullets(doc, [
        'The buttons across the top of your roster are your **class periods**. Tap one to see '
        'just that class. The number is how many students are in it.',
        'Use the **search box** to find one student fast.',
        '**Remove** takes a student out of **that one class**. If you have them in two periods, '
        'the other class isn\u2019t touched.',
        'Whole class in the wrong period? Tap that period, then **Clear this period off my '
        'roster**, then upload the file again under the right period.',
    ])
    callout(doc, 'Good to know:',
            'Clearing or removing **never deletes a student or their QR code.** Their printed '
            'card keeps working and their history is kept. You\u2019re only changing your class '
            'list.')

    # ── 3 ──────────────────────────────────────────────
    h(doc, '3.  Writing a pass \u2014 on your computer')
    steps(doc, [
        'Your roster is the front page. Find the student.',
        'Tap **Log OUT** on their row.',
        'Pick where they\u2019re going.',
        'Tap the big blue **Log Pass OUT**.',
    ])
    sub(doc, 'The pass types')
    bullets(doc, [
        '**Restroom**',
        '**Nurse** \u2014 opens boxes to check what\u2019s wrong and what you already tried '
        '(water, bandaid, head down\u2026). Check what applies; you don\u2019t have to fill in '
        'everything.',
        '**Office**',
        '**Student Services** \u2014 pick who they\u2019re seeing, or pick Other and type a name.',
        '**Late Departure** \u2014 type who\u2019s releasing them.',
        '**Going to Another Teacher** \u2014 type which teacher.',
    ])
    bullets(doc, [
        'Their row now says **OUT** and starts counting minutes.',
        'A banner at the top tells you how many of your students are out right now.',
    ])

    # ── 4 ──────────────────────────────────────────────
    h(doc, '4.  Writing a pass \u2014 on your phone')
    doc.add_paragraph('Same thing, but you scan the student\u2019s QR card instead of hunting for '
                      'their name.')
    steps(doc, [
        '**Sign in to Muster on your phone first.** Do this once, before you ever scan.',
        'Open your phone\u2019s **regular Camera app** and point it at the QR card.',
        'Tap the link that pops up.',
        'Pick the pass type, tap **Log Pass OUT**. Done.',
    ])
    callout(doc, 'Two things that trip people up:', '', RGBColor(0xB5, 0x84, 0x00))
    bullets(doc, [
        '**Use the built-in Camera app.** Downloaded \u201cQR scanner\u201d apps open the page in '
        'their own little browser, which can\u2019t sign you in to Google \u2014 you get stuck on '
        'a login screen that won\u2019t work.',
        '**Sign in before you scan**, not after. Scan first and you may just get bounced to a '
        'sign-in page.',
    ], indent=0.4)

    # ── 5 ──────────────────────────────────────────────
    h(doc, '5.  Logging them back in')
    bullets(doc, [
        '**On the computer:** their row now shows a green **Log IN**. Tap it, then tap '
        '**Log Back In**.',
        '**On your phone:** scan the same card again. Muster already knows they\u2019re out, so it '
        'just offers the green **Log Back In** button.',
    ])
    callout(doc, 'You never pick in or out.',
            '**Scanning the same card does whichever one makes sense** \u2014 out if they\u2019re '
            'in the room, in if they\u2019re already out.')
    bullets(doc, [
        'Forgot to log someone back in? Just do it when you notice. The minutes will look long; '
        'that\u2019s fine, nobody\u2019s in trouble.',
    ])

    # ── 6 ──────────────────────────────────────────────
    h(doc, '6.  Fire drill or lockdown')
    doc.add_paragraph('This is the part worth practicing before you need it.')
    sub(doc, 'Students on your roster')
    steps(doc, [
        'Tap **My Roll Call** at the top.',
        'It already shows the class you have right now.',
        'Tap **Secure** beside each student who is with you.',
        'Their row turns green and the count goes up. The office sees it instantly.',
    ])
    bullets(doc, [
        'Keep tapping \u2014 the page doesn\u2019t reload and won\u2019t lose your place.',
        'If a button turns yellow and says **Retry**, tap it again. That\u2019s a weak signal, not '
        'a mistake you made.',
    ])
    sub(doc, 'A student who isn\u2019t yours')
    steps(doc, [
        'Scan their QR card with your phone.',
        'Tap the big red **Mark SECURE with me** at the top of the page.',
    ])
    callout(doc, 'On purpose:',
            'you can scan **any** student in the school, not just your own \u2014 during a drill '
            'kids end up in the wrong room.')

    # ── 7 ──────────────────────────────────────────────
    h(doc, '7.  Printing student cards')
    doc.add_paragraph('Every student needs their QR code on something. Tap Print at the top.')
    bullets(doc, [
        '**Lanyard cards** \u2014 big cards, one per student. Laminate them.',
        '**ID Stickers** \u2014 small, about 32 to a page. Stick them on student IDs.',
    ])
    steps(doc, [
        'Tap a class period first if you only want that class.',
        '**Wait for every QR square to finish loading.** Print too fast and you get blank boxes.',
        'Tap **Print**.',
    ])
    bullets(doc, [
        'Print once and you\u2019re done \u2014 a student\u2019s code never changes, all year.',
        'Lost card? Reprint that class; it comes out the same.',
    ])

    # ── 8 ──────────────────────────────────────────────
    h(doc, '8.  Warnings you might see')
    doc.add_paragraph('Muster shows these for you to use your judgment on. It never blocks a '
                      'student on its own.')
    bullets(doc, [
        '**Used a pass in the last hour** \u2014 they just went somewhere.',
        '**3 passes today** (red number by their name) \u2014 going a lot today.',
        '**2nd day in a row** \u2014 leaving the same class two days running. Sometimes a pattern, '
        'sometimes a stomach bug.',
        '**Too many students in the hallway** \u2014 the school-wide restroom count is at its cap. '
        'Wait for it to drop before sending another.',
        '**Cannot leave without an adult escort** \u2014 an administrator flagged this student, '
        'and there may be a note saying why. Pass options disappear. Don\u2019t send them; call '
        'the office.',
    ])

    # ── 9 ──────────────────────────────────────────────
    h(doc, '9.  When something looks wrong')
    bullets(doc, [
        '**My roster is empty.** You haven\u2019t added students yet \u2014 see section 2.',
        '**A card won\u2019t scan.** Check the QR is clean and glare-free, then look the student up '
        'on your roster and tap **Scan** instead. Same screen, no camera needed.',
        '**Scanning sends me to a login page over and over.** You\u2019re in a QR-scanner app\u2019s '
        'browser. Close it, open Muster in Safari or Chrome, sign in, then scan with the normal '
        'Camera app.',
        '**A student shows OUT but they\u2019re sitting right here.** Log them back in \u2014 '
        'someone forgot. No harm done.',
        '**I sent someone out from the wrong period.** Leave it. The time is recorded correctly; '
        'only the period label is off.',
        '**A student is missing from My Roll Call.** They\u2019re probably filed under a different '
        'period. Tap the other period buttons at the top.',
        '**A name is spelled wrong, or a student left the school.** Ask Ms. Howe \u2014 those are '
        'admin jobs.',
    ])
    callout(doc, 'Still stuck?',
            'Tell Ms. Howe two things: **what you tapped, and what it said back.** That\u2019s '
            'usually enough to sort it out in a minute.')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run('Muster (by Howe)  \u00b7  the same guide lives in the app under \u201cHelp\u201d')
    r.font.size, r.font.color.rgb = Pt(9), GREY

    doc.save(OUT)
    print(f'Wrote {OUT}')
    return OUT


if __name__ == '__main__':
    build()
